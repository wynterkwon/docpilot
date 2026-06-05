from __future__ import annotations

from pathlib import Path

from docpilot.exceptions import IngestionError
from docpilot.ingestion.models import IngestedDocument


def ingest(path: str | Path) -> IngestedDocument:
    path = Path(path)

    if not path.exists():
        raise IngestionError("File not found", detail=str(path))

    if path.suffix.lower() != ".docx":
        raise IngestionError(f"Expected .docx, got '{path.suffix}'", detail=str(path))

    try:
        from docx import Document
    except ImportError as e:
        raise IngestionError("python-docx is required: pip install \"docpilot[docx]\"") from e

    try:
        doc = Document(str(path))
    except Exception as e:
        raise IngestionError("Failed to open DOCX", detail=str(e)) from e

    parts: list[str] = []

    for block in _iter_blocks(doc):
        if block:
            parts.append(block)

    content = "\n\n".join(parts)

    return IngestedDocument(
        source=path,
        content=content,
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        metadata={
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "size_bytes": path.stat().st_size,
        },
    )


def _iter_blocks(doc):
    """Yield text blocks from paragraphs and tables in document order."""
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body = doc.element.body
    for child in body:
        if child.tag == qn("w:p"):
            para = Paragraph(child, doc)
            text = para.text.strip()
            if text:
                style = para.style.name if para.style else ""
                if style.startswith("Heading"):
                    yield f"[{text}]"
                else:
                    yield text
        elif child.tag == qn("w:tbl"):
            table = Table(child, doc)
            yield _table_to_text(table)


def _table_to_text(table) -> str:
    rows: list[str] = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(" | ".join(cells))
    return "\n".join(rows)
