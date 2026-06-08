from __future__ import annotations

import copy
from pathlib import Path

from docpilot.builder.base import BaseBuilder, PLACEHOLDER_RE
from docpilot.exceptions import BuilderError


class DocxBuilder(BaseBuilder):
    def build(
        self,
        template: str | Path,
        sections: dict[str, str],
        output: str | Path,
    ) -> Path:
        template, output = self._validate_paths(template, output)

        if template.suffix.lower() != ".docx":
            raise BuilderError(f"Expected .docx template, got '{template.suffix}'")

        try:
            import docx
        except ImportError as e:
            raise BuilderError("python-docx is required: pip install python-docx") from e

        try:
            doc = docx.Document(str(template))
        except Exception as e:
            raise BuilderError("Failed to open DOCX template", detail=str(e)) from e

        # list() snapshot — new paragraphs inserted during iteration are skipped
        for para in list(doc.paragraphs):
            _replace_in_paragraph(para, sections)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in list(cell.paragraphs):
                        _replace_in_paragraph(para, sections)

        try:
            doc.save(str(output))
        except Exception as e:
            raise BuilderError("Failed to save DOCX", detail=str(e)) from e

        return output


def _replace_in_paragraph(para, sections: dict[str, str]) -> None:
    full_text = "".join(run.text for run in para.runs)
    match = PLACEHOLDER_RE.search(full_text)
    if not match:
        return

    key = match.group(1)
    if key not in sections:
        return

    replaced = PLACEHOLDER_RE.sub(sections[key], full_text, count=1)
    lines = replaced.split("\n")

    # Set first line into the original paragraph (preserves run formatting)
    if para.runs:
        para.runs[0].text = lines[0]
        for run in para.runs[1:]:
            run.text = ""

    if len(lines) == 1:
        return

    # Multi-line: clone this paragraph's XML for each subsequent line
    from docx.oxml.ns import qn

    p_elem = para._element
    parent = p_elem.getparent()
    if parent is None:
        # No parent — fall back to joining with spaces
        if para.runs:
            para.runs[0].text = " ".join(lines)
        return

    idx = list(parent).index(p_elem)
    for i, line in enumerate(lines[1:], 1):
        new_p = copy.deepcopy(p_elem)
        runs = new_p.findall(qn("w:r"))
        if runs:
            t_list = runs[0].findall(qn("w:t"))
            if t_list:
                t_list[0].text = line
                for extra_t in t_list[1:]:
                    runs[0].remove(extra_t)
            for extra_r in runs[1:]:
                new_p.remove(extra_r)
        parent.insert(idx + i, new_p)
