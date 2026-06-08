"""DOCX 템플릿에서 플레이스홀더별 스타일 컨텍스트를 추출한다."""
from __future__ import annotations

from pathlib import Path

from docpilot.builder.base import PLACEHOLDER_RE

# DXA(twip) → mm: 1 twip = 1/1440 inch = 25.4/1440 mm
_DXA_PER_MM = 1440 / 25.4


def extract_style_hints(template_path: Path) -> dict[str, str]:
    """DOCX 템플릿을 분석해 플레이스홀더별 자연어 스타일 힌트를 반환한다.

    Returns:
        {"섹션명": "10pt 맑은 고딕, JUSTIFY, 표 셀 너비 60mm"} 형식의 dict
        스타일 정보를 추출할 수 없으면 빈 dict 반환
    """
    if template_path.suffix.lower() != ".docx":
        return {}

    try:
        import docx as python_docx
    except ImportError:
        return {}

    try:
        doc = python_docx.Document(str(template_path))
    except Exception:
        return {}

    hints: dict[str, str] = {}

    for para in doc.paragraphs:
        _process_para(para, hints)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _process_para(para, hints)

    return hints


def _process_para(para, hints: dict[str, str]) -> None:
    full_text = "".join(run.text for run in para.runs)
    match = PLACEHOLDER_RE.search(full_text)
    if not match:
        return

    key = match.group(1)
    if key in hints:
        return

    # Run that directly contains the placeholder text (fallback: first run)
    target_run = next(
        (r for r in para.runs if PLACEHOLDER_RE.search(r.text)),
        para.runs[0] if para.runs else None,
    )

    parts: list[str] = []

    if target_run is not None:
        pt = _get_font_size_pt(target_run, para)
        if pt:
            font = _get_east_asia_font(target_run) or _get_font_name(target_run, para)
            pt_str = f"{pt:g}pt"
            parts.append(f"{pt_str} {font}".strip() if font else pt_str)

        if _is_bold(target_run, para):
            parts.append("볼드")
        if _is_italic(target_run, para):
            parts.append("이탤릭")

    align = _get_alignment(para)
    if align:
        parts.append(align)

    width_mm = _get_cell_width_mm(para)
    if width_mm:
        parts.append(f"표 셀 너비 {width_mm:.0f}mm")

    hints[key] = ", ".join(parts) if parts else ""


# ---------------------------------------------------------------------------
# Style resolution helpers
# ---------------------------------------------------------------------------

def _get_font_size_pt(run, para) -> float | None:
    if run.font.size is not None:
        return run.font.size.pt
    if run.style and run.style.font.size is not None:
        return run.style.font.size.pt
    style = para.style
    while style is not None:
        if style.font.size is not None:
            return style.font.size.pt
        style = style.base_style
    return None


def _get_font_name(run, para) -> str:
    if run.font.name:
        return run.font.name
    if run.style and run.style.font.name:
        return run.style.font.name
    style = para.style
    while style is not None:
        if style.font.name:
            return style.font.name
        style = style.base_style
    return ""


def _get_east_asia_font(run) -> str:
    """w:rFonts의 eastAsia 속성(한글 글꼴)을 반환한다."""
    try:
        from docx.oxml.ns import qn
        r_pr = run._element.find(qn("w:rPr"))
        if r_pr is None:
            return ""
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            return ""
        return r_fonts.get(qn("w:eastAsia")) or ""
    except Exception:
        return ""


def _is_bold(run, para) -> bool:
    if run.bold is not None:
        return run.bold
    if run.style and run.style.font.bold is not None:
        return run.style.font.bold
    style = para.style
    while style is not None:
        if style.font.bold is not None:
            return style.font.bold
        style = style.base_style
    return False


def _is_italic(run, para) -> bool:
    if run.italic is not None:
        return run.italic
    if run.style and run.style.font.italic is not None:
        return run.style.font.italic
    style = para.style
    while style is not None:
        if style.font.italic is not None:
            return style.font.italic
        style = style.base_style
    return False


def _get_alignment(para) -> str:
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return ""

    _MAP = {
        WD_ALIGN_PARAGRAPH.CENTER: "CENTER",
        WD_ALIGN_PARAGRAPH.LEFT: "LEFT",
        WD_ALIGN_PARAGRAPH.RIGHT: "RIGHT",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY",
        WD_ALIGN_PARAGRAPH.JUSTIFY_LOW: "JUSTIFY",
        WD_ALIGN_PARAGRAPH.DISTRIBUTE: "JUSTIFY",
    }

    align = para.alignment
    if align is None:
        style = para.style
        while style is not None:
            pf = getattr(style, "paragraph_format", None)
            if pf is not None and pf.alignment is not None:
                align = pf.alignment
                break
            style = style.base_style

    return _MAP.get(align, "") if align is not None else ""


def _get_cell_width_mm(para) -> float | None:
    """para가 표 셀 안에 있으면 셀 너비(mm)를 반환한다."""
    try:
        from docx.oxml.ns import qn
        elem = para._element
        parent = elem.getparent()
        while parent is not None:
            if parent.tag == qn("w:tc"):
                tc_pr = parent.find(qn("w:tcPr"))
                if tc_pr is None:
                    return None
                tc_w = tc_pr.find(qn("w:tcW"))
                if tc_w is None:
                    return None
                w_type = tc_w.get(qn("w:type"))
                w_val = tc_w.get(qn("w:w"))
                if w_val and w_type in ("dxa", None):
                    return int(w_val) / _DXA_PER_MM
                return None
            parent = parent.getparent()
    except Exception:
        pass
    return None
