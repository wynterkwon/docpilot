from __future__ import annotations

import zipfile
from pathlib import Path

import pytest

_HP_NS = "http://www.hancom.co.kr/hwpml/2012/paragraph"
_HH_NS = "http://www.hancom.co.kr/hwpml/2011/head"


# ---------------------------------------------------------------------------
# HWPX fixture helpers
# ---------------------------------------------------------------------------

def _hwpx_basic(path: Path) -> Path:
    """20pt 바탕 볼드 JUSTIFY {{서론}} 문단이 있는 HWPX."""
    header = f"""<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="{_HH_NS}">
  <hh:charPr id="0" height="2000">
    <hh:fontRef hangul="바탕"/>
    <hh:bold/>
  </hh:charPr>
  <hh:paraPr id="0" align="JUSTIFY"/>
</hh:head>""".encode()

    content = f"""<?xml version="1.0" encoding="UTF-8"?>
<hml xmlns:hp="{_HP_NS}">
  <hp:body><hp:sec>
    <hp:p id="1" paraPrIDRef="0">
      <hp:run charPrIDRef="0"><hp:t>{{{{서론}}}}</hp:t></hp:run>
    </hp:p>
  </hp:sec></hp:body>
</hml>""".encode()

    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("mimetype", "application/hwp+zip", compress_type=zipfile.ZIP_STORED)
        zf.writestr("Contents/header.xml", header)
        zf.writestr("Contents/content.hml", content)
    return path


def _hwpx_table(path: Path) -> Path:
    """셀 너비 34015 HWPUNIT(≈120mm) 표 안에 {{표 내용}} 문단이 있는 HWPX."""
    header = f"""<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="{_HH_NS}">
  <hh:charPr id="0" height="1000"/>
  <hh:paraPr id="0" align="LEFT"/>
</hh:head>""".encode()

    content = f"""<?xml version="1.0" encoding="UTF-8"?>
<hml xmlns:hp="{_HP_NS}">
  <hp:body><hp:sec>
    <hp:tbl>
      <hp:tr>
        <hp:tc>
          <hp:cellSz width="34015"/>
          <hp:p id="1" paraPrIDRef="0">
            <hp:run charPrIDRef="0"><hp:t>{{{{표 내용}}}}</hp:t></hp:run>
          </hp:p>
        </hp:tc>
      </hp:tr>
    </hp:tbl>
  </hp:sec></hp:body>
</hml>""".encode()

    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("mimetype", "application/hwp+zip", compress_type=zipfile.ZIP_STORED)
        zf.writestr("Contents/header.xml", header)
        zf.writestr("Contents/content.hml", content)
    return path


# ---------------------------------------------------------------------------
# HWPX analyzer tests
# ---------------------------------------------------------------------------

class TestHwpxAnalyzer:
    def test_font_size_and_bold(self, tmp_path: Path):
        from docpilot.builder.hwpx_analyzer import extract_style_hints
        hints = extract_style_hints(_hwpx_basic(tmp_path / "t.hwpx"))
        assert "서론" in hints
        assert "20pt" in hints["서론"]
        assert "바탕" in hints["서론"]
        assert "볼드" in hints["서론"]

    def test_alignment(self, tmp_path: Path):
        from docpilot.builder.hwpx_analyzer import extract_style_hints
        hints = extract_style_hints(_hwpx_basic(tmp_path / "t.hwpx"))
        assert "JUSTIFY" in hints["서론"]

    def test_cell_width(self, tmp_path: Path):
        from docpilot.builder.hwpx_analyzer import extract_style_hints
        hints = extract_style_hints(_hwpx_table(tmp_path / "t.hwpx"))
        assert "표 내용" in hints
        assert "표 셀 너비" in hints["표 내용"]
        assert "120mm" in hints["표 내용"]  # 34015 / 283.46 ≈ 120mm

    def test_wrong_extension_returns_empty(self, tmp_path: Path):
        from docpilot.builder.hwpx_analyzer import extract_style_hints
        bad = tmp_path / "t.docx"
        bad.write_bytes(b"fake")
        assert extract_style_hints(bad) == {}

    def test_no_placeholder_returns_empty(self, tmp_path: Path):
        from docpilot.builder.hwpx_analyzer import extract_style_hints
        content = f"""<?xml version="1.0" encoding="UTF-8"?>
<hml xmlns:hp="{_HP_NS}">
  <hp:body><hp:sec>
    <hp:p><hp:run><hp:t>플레이스홀더 없음</hp:t></hp:run></hp:p>
  </hp:sec></hp:body>
</hml>""".encode()
        path = tmp_path / "no_ph.hwpx"
        with zipfile.ZipFile(path, "w") as zf:
            zf.writestr("mimetype", "application/hwp+zip", compress_type=zipfile.ZIP_STORED)
            zf.writestr("Contents/content.hml", content)
        assert extract_style_hints(path) == {}


# ---------------------------------------------------------------------------
# DOCX analyzer tests
# ---------------------------------------------------------------------------

class TestDocxAnalyzer:
    def test_font_size_and_bold(self, tmp_path: Path):
        pytest.importorskip("docx")
        from docx import Document
        from docx.shared import Pt
        from docpilot.builder.docx_analyzer import extract_style_hints

        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("{{서론}}")
        run.font.size = Pt(14)
        run.bold = True
        path = tmp_path / "t.docx"
        doc.save(str(path))

        hints = extract_style_hints(path)
        assert "서론" in hints
        assert "14pt" in hints["서론"]
        assert "볼드" in hints["서론"]

    def test_alignment(self, tmp_path: Path):
        pytest.importorskip("docx")
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docpilot.builder.docx_analyzer import extract_style_hints

        doc = Document()
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run("{{결론}}")
        path = tmp_path / "t.docx"
        doc.save(str(path))

        hints = extract_style_hints(path)
        assert "결론" in hints
        assert "CENTER" in hints["결론"]

    def test_cell_width(self, tmp_path: Path):
        pytest.importorskip("docx")
        from docx import Document
        from docx.shared import Mm
        from docpilot.builder.docx_analyzer import extract_style_hints

        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        cell.width = Mm(60)
        cell.paragraphs[0].add_run("{{표 내용}}")
        path = tmp_path / "t.docx"
        doc.save(str(path))

        hints = extract_style_hints(path)
        assert "표 내용" in hints
        assert "표 셀 너비" in hints["표 내용"]

    def test_wrong_extension_returns_empty(self, tmp_path: Path):
        from docpilot.builder.docx_analyzer import extract_style_hints
        bad = tmp_path / "t.hwpx"
        bad.write_bytes(b"fake")
        assert extract_style_hints(bad) == {}

    def test_no_placeholder_returns_empty(self, tmp_path: Path):
        pytest.importorskip("docx")
        from docx import Document
        from docpilot.builder.docx_analyzer import extract_style_hints

        doc = Document()
        doc.add_paragraph("플레이스홀더 없음")
        path = tmp_path / "no_ph.docx"
        doc.save(str(path))
        assert extract_style_hints(path) == {}
