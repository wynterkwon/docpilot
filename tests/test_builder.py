from __future__ import annotations

from pathlib import Path

import pytest

from docpilot.builder.hwpx_builder import HwpxBuilder
from docpilot.exceptions import BuilderError

SECTIONS = {"서론": "서론 내용입니다.", "결론": "결론 내용입니다."}


class TestHwpxBuilder:
    def test_build_produces_file(self, hwpx_template: Path, tmp_path: Path):
        output = tmp_path / "output.hwpx"
        builder = HwpxBuilder()
        result = builder.build(hwpx_template, SECTIONS, output)
        assert result == output
        assert output.exists()
        assert output.stat().st_size > 0

    def test_placeholders_replaced(self, hwpx_template: Path, tmp_path: Path):
        import zipfile
        output = tmp_path / "output.hwpx"
        HwpxBuilder().build(hwpx_template, SECTIONS, output)

        with zipfile.ZipFile(output, "r") as zf:
            candidates = [n for n in zf.namelist() if n.endswith("content.hml")]
            content = zf.read(candidates[0]).decode("utf-8")

        assert "서론 내용입니다." in content
        assert "{{서론}}" not in content

    def test_template_not_found_raises(self, tmp_path: Path):
        with pytest.raises(BuilderError, match="not found"):
            HwpxBuilder().build(tmp_path / "missing.hwpx", SECTIONS, tmp_path / "out.hwpx")

    def test_wrong_extension_raises(self, tmp_path: Path):
        bad = tmp_path / "template.docx"
        bad.write_bytes(b"fake")
        with pytest.raises(BuilderError, match="Expected .hwpx"):
            HwpxBuilder().build(bad, SECTIONS, tmp_path / "out.hwpx")


class TestDocxBuilder:
    def test_build_produces_file(self, tmp_path: Path):
        docx = pytest.importorskip("docx")
        from docx import Document
        from docpilot.builder.docx_builder import DocxBuilder

        template = tmp_path / "template.docx"
        doc = Document()
        doc.add_paragraph("{{서론}}")
        doc.add_paragraph("{{결론}}")
        doc.save(str(template))

        output = tmp_path / "output.docx"
        DocxBuilder().build(template, SECTIONS, output)

        assert output.exists()
        result_doc = Document(str(output))
        texts = [p.text for p in result_doc.paragraphs]
        assert "서론 내용입니다." in texts
        assert "결론 내용입니다." in texts
